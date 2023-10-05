from docx import Document
from docx.enum.text import WD_COLOR
from docxcompose.composer import Composer
import pandas as pd
import glob
import re
import os

# Import keywords from the Excel file and turn them into a list
df = pd.read_excel('keywords.xlsx')
list_of_words = df['keywords'].tolist()

# Set up regex
patterns = [r'\b' + word + r'\b' for word in list_of_words]
re_highlight = re.compile('(' + '|'.join(p for p in patterns) + ')+', re.IGNORECASE)

# Define a function to extract the file name without extension
def extract_file_name(file):
    file_name = os.path.basename(file)
    file_name_without_extension = os.path.splitext(file_name)[0]
    return file_name_without_extension

# Define a function to process all *.docx files in a folder
def process_docx_files(folder_path):
    docx_files = glob.glob(folder_path + '\\*.docx')
    for file in docx_files:
        doc = Document(file)
        title = extract_file_name(file)  # Extract the file name without extension
        modified_doc = Document()  # Create a new document to store modified paragraphs
        modified_doc.add_heading(title, level=1)  # Add the title as Header 1 to the modified document
        for para in doc.paragraphs:
            text = para.text
            if len(re_highlight.findall(text)) > 0:
                matches = re_highlight.finditer(text)
                p3 = 0
                highlighted_para = modified_doc.add_paragraph()  # Add a new paragraph to the modified document
                for match in matches:
                    p1 = p3
                    p2, p3 = match.span()
                    highlighted_para.add_run(text[p1:p2])
                    run = highlighted_para.add_run(text[p2:p3])
                    run.font.highlight_color = WD_COLOR.YELLOW
                    highlighted_para.add_run(text[p3:])
        if modified_doc.paragraphs:  # Only save the modified document if it contains highlighted paragraphs
            modified_doc.save(file)

# Provide the folder path where the algorithm will iterate over all *.docx files
folder_path = '.\\sample.docx'
process_docx_files(folder_path)

