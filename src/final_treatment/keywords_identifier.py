from docx import Document
from docx.enum.text import WD_COLOR
import pandas as pd
import glob
import re

df = pd.read_excel('keywords.xlsx')
list_of_words = df['keywords'].tolist()

# Setup regex
patterns = [r'\b' + word + r'\b' for word in list_of_words]
re_highlight = re.compile('(' + '|'.join(p for p in patterns) + ')+', re.IGNORECASE)

folder_path = '.\\sample.docx' 

docx_files = glob.glob(folder_path + '\\*.docx')

for file in docx_files:
    doc = Document(file)
    for para in doc.paragraphs:
        text = para.text
        if len(re_highlight.findall(text)) > 0:
            matches = re_highlight.finditer(text)
            para.text = ''
            p3 = 0
            for match in matches:
                p1 = p3
                p2, p3 = match.span()
                para.add_run(text[p1:p2])
                run = para.add_run(text[p2:p3])
                run.font.highlight_color = WD_COLOR.YELLOW
                para.add_run(text[p3:])
    doc.save(file)