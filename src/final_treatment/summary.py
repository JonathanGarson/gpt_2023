import os
import glob
import json
from langchain.vectorstores import Chroma
from langchain.embeddings import OpenAIEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.chat_models import ChatOpenAI
from langchain.chains import RetrievalQA
from langchain.document_loaders import TextLoader
from langchain.document_loaders import DirectoryLoader
from langchain.vectorstores import Chroma
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.document_loaders import Docx2txtLoader


#Kor!
from kor.extraction import create_extraction_chain
from kor.nodes import Object, Text, Number
from pydantic import BaseModel, Field, validator
from kor import extract_from_documents, from_pydantic, create_extraction_chain

# LangChain Models
from langchain.chat_models import ChatOpenAI
from langchain.embeddings.openai import OpenAIEmbeddings

# Standard Helpers
import pandas as pd
import glob
from datetime import datetime
import json
import os

#from openai
import openai

#from typing import Optional
from typing import Optional

#from numpy
import numpy as np

#from pydocx
from docx import Document

#from shutil
import shutil

os.environ["OPENAI_API_KEY"] = "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"

openai_api_key = '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'

# Step 1: Define Kor objects and models
augmentation_gen_ouv = Object(
    id="augmentation_generale_ouv",
    description="augmentation générale de salaire/salaire de base pour les ouvriers et les employés, en pourcentage ou en euros",
    attributes=[
        Number(id="ouvrier", description="augmentation générale de salaire/salaire de base pour un employé/ouvrier, en pourcentage ou en euros")
    ],
    examples=[
        (
            """5% au titre de l’Augmentation Générale, à l’ensemble des collaborateurs ouvriers, employés. Les techniciens \
            et agents de maitrise percevant en 2022 la PPCU (prime performance collective usine), en contrat CDD/CDI, \ 
            justifiant d’une ancienneté supérieure ou égale à 6 mois au 31 décembre 2022 bénéficeriont d'une augmentation de 4%. Les cadres bénéfieceront d'une augmentation des salaires de base \
            de 3%""",
            [
                {"ouvrier": "5%"}
            ]),
            
           ( """Pour les non-cadres : Une augmentation générale de 4,5 % applicable en novembre 2022\
            avec une augmentation minimale de 100 euros brut mensuel. Cette augmentation minimale sera versée en janvier 2023.""",
            [
                {"ouvrier": "4,5"}
            ]
        )
    ]
)

augmentation_gen_int = Object(
    id="augmentation_generale_int",
    description="augmentation générale de salaire/salaire de base pour les proféssions intermédiaires et agent de maîtrise, techniciens, en pourcentage ou en euros",
    attributes=[
        Number(id="intermediaires", description="augmentation générale de salaire/salaire de base pour une profession intermédiaire, technicien, agent de maitrise, en pourcentage ou en euros")
    ],
    examples=[
        (
            """5% au titre de l’Augmentation Générale, à l’ensemble des collaborateurs ouvriers, employés. Les techniciens \
            et agents de maitrise percevant en 2022 la PPCU (prime performance collective usine), en contrat CDD/CDI, \ 
            justifiant d’une ancienneté supérieure ou égale à 6 mois au 31 décembre 2022 bénéficeriont d'une augmentation de 4%. Les cadres bénéfieceront d'une augmentation des salaires de base \
            de 3%""",
            [
                {"ouvrier": "5%"},
                {"intermédiaires": "4%"},
                {"cadres": "3%"}
            ]),
            
            ("""Pour les non-cadres : Une augmentation générale de 4,5 % applicable en novembre 2022\
            avec une augmentation minimale de 100 euros brut mensuel. Cette augmentation minimale sera versée en janvier 2023.""",
            [
                {"intermédiaires": "4,5"}
            ],
        )
    ]
)

augmentation_gen_cad = Object(
    id="augmentation_generale_csp",
    description="augmentation générale de salaire/salaire de base pour les cadres et ingénieurs, en pourcentage ou en euros",
    attributes=[
        Number(id = "cadres", description="augmentation générale de salaire/salaire de base pour un cadre,ingénieur, en pourcentage ou en euros")
    ],
    examples=[
        (
            """5% au titre de l’Augmentation Générale, à l’ensemble des collaborateurs ouvriers, employés, techniciens \
            et agents de maitrise percevant en 2022 la PPCU (prime performance collective usine), en contrat CDD/CDI, \ 
            justifiant d’une ancienneté supérieure ou égale à 6 mois au 31 décembre 2022. Les cadres bénéfiecerons d'une augmentation de 4%""",
            [
                {"cadres": "3%"}
            ]),
            
            ("""Pour les non-cadres : Une augmentation générale de 4,5 % applicable en novembre 2022\
            avec une augmentation minimale de 100 euros brut mensuel. Cette augmentation minimale sera versée en janvier 2023.""",
            [
                {"cadres": "0"}
            ]
        )
    ]
)

augmentation_gen = Object(
    id="augmentation_generale",
    description="Information sur l'augmentation de la masse salariale",
    examples=[
        (
            """A la suite de la demande de la Délégation syndicale CFDT, la Direction accepte de réserver une enveloppe de 5% de la Masse salariale. 5% au titre de l’Augmentation Générale,\ 
            à l’ensemble des collaborateurs ouvriers, employés, techniciens \
            et agents de maitrise percevant en 2022 la PPCU (prime performance collective usine), en contrat CDD/CDI, justifiant d’une ancienneté supérieure ou égale à 6 mois au 31 décembre 2022""",
            [
                {"augmentation générale": "1%"}
            ]),
            (""" Concernant les augmentations individuelles, la Direction consent l’attribution d’une enveloppe de 5% de la masse salariale.""",
            [
                {"augmentation générale" : "5%"}
            ]
        )
    ],
    attributes=[
        Number(
            id="augmentation_generale_tous_salaries",
            description="Augmentation générale accordée à tous les salariés indépendemment de leur catégorie socio-professionnelle (ouvrier, intermédiaire, cadres), en pourcentage ou en euros",
        )
    ]
)

augmentation_ind_ouv = Object(
    id="augmentation_ind_ouv",
    description="augmentation individuelle de salaire ou mérite pour les ouvriers et les employés, en pourcentage ou en euros",
    attributes=[
        Number(id="ouvrier", description="augmentation individuelle de salaire ou mérite pour un employé/ouvrier, en pourcentage ou en euros")
    ],
    examples=[
        (
            """Mise en œuvre d’une révision salariale annuelle pérenne de 3% attribuée aux collaborateurs avec un salaire annuel brut de base inférieur ou égal à 50 000,00 euros,\ 
            de  2% pour les collaborateurs avec un salaire annuel brut de base strictement supérieur à 50 000,00 euros et inférieur ou égal à 70 000,00 euros \ 
            et de 1% pour les collaborateurs avec un salaire annuel brut de base strictement supérieur à 70 000,00 euros et inférieur ou égal à 90 000,00 euros.""",
            [
                {"ouvrier": "3"}
            ]),
            ("""Augmentation individuelle pour les ouvriers et employés:
            Une enveloppe de 2% de la masse salariale du personnel Ouvrier et employé sera destinée à des augmentations individuelles. 
            """,
            [
                {"ouvrier":"2"}
            ]
        )
    ]
)

augmentation_ind_int = Object(
    id="augmentation_ind_int",
    description="augmentation individuelle de salaire ou mérite pour les professions intermédiaires, techniniciens, en pourcentage ou en euros",
    attributes=[
        
        Number(id="intermediaires", description="augmentation individuelle de salaire ou mérite pour une profession intermédiaire, technicien, agent de maitrise, en pourcentage ou en euros")
    ],
    examples=[
        (
            """Mise en œuvre d’une révision salariale annuelle pérenne de 3% attribuée aux collaborateurs avec un salaire annuel brut de base inférieur ou égal à 50 000,00 euros,\ 
            de  2% pour les collaborateurs avec un salaire annuel brut de base strictement supérieur à 50 000,00 euros et inférieur ou égal à 70 000,00 euros \ 
            et de 1% pour les collaborateurs avec un salaire annuel brut de base strictement supérieur à 70 000,00 euros et inférieur ou égal à 90 000,00 euros.""",
            [
                {"intermédiaires": "2%"},
            ]),
            ("""Pour les catégories Agents de maîtrise et Cadres  : 

            Sans remettre en cause le principe d’une enveloppe globale destinée aux augmentations individuelles, conformément aux procédures applicables au sein de la société il est convenu, par anticipation sur ladite enveloppe,\
            d’une augmentation pour l’ensemble des salariés desdites catégories, de 3.5 % à valoir sur les augmentations individuelles à intervenir en février 2023.""",
            [
                {"intermédiaires":" 3.5%"}
            ]

        )
    ]
)

augmentation_ind_cad = Object(
    id="augmentation_ind_cad",
    description="augmentation individuelle de salaire ou mérite pour les cadres et ingénieurs, en pourcentage ou en euros",
    attributes=[
        Number(id = "cadres", description="augmentation individuelle de salaire ou mérite pour un cadre,ingénieur, en pourcentage ou en euros")
    ],
    examples=[
        (
            """Mise en œuvre d’une révision salariale annuelle pérenne de 3% attribuée aux collaborateurs avec un salaire annuel brut de base inférieur ou égal à 50 000,00 euros,\ 
            de  2% pour les collaborateurs avec un salaire annuel brut de base strictement supérieur à 50 000,00 euros et inférieur ou égal à 70 000,00 euros \ 
            et de 1% pour les collaborateurs avec un salaire annuel brut de base strictement supérieur à 70 000,00 euros et inférieur ou égal à 90 000,00 euros.""",
            [
                {"cadres": "1%"}
            ]),
            ("""Pour les catégories Agents de maîtrise et Cadres  : 

            Sans remettre en cause le principe d’une enveloppe globale destinée aux augmentations individuelles, conformément aux procédures applicables au sein de la société il est convenu, par anticipation sur ladite enveloppe,\
            d’une augmentation pour l’ensemble des salariés desdites catégories, de 3.5 % à valoir sur les augmentations individuelles à intervenir en février 2023.""",
            [
                {"cadres":" 3.5%"}
            ]

        )
    ]
)


augmentation_ind = Object(
    id="augmentation_ind",
    description="Information sur les augmentations individuelles ou au mérite",
    examples=[
        (
           """Mise en œuvre d’une révision salariale annuelle pérenne de 3% attribuée aux collaborateurs avec un salaire annuel brut de base inférieur ou égal à 50 000,00 euros,\ 
            de  2% pour les collaborateurs avec un salaire annuel brut de base strictement supérieur à 50 000,00 euros et inférieur ou égal à 70 000,00 euros \ 
            et de 1% pour les collaborateurs avec un salaire annuel brut de base strictement supérieur à 70 000,00 euros et inférieur ou égal à 90 000,00 euros. \
            La direction dédit une enveloppe de 1% qui sera réparti à la performance et au mérite à la discretion des managers""",
            [
                {"augmentation individuelle": "1%"}
            ], 
        )
    ],
    attributes=[
        Number(
            id="augmentation_individuelle_tous_salaries",
            description="Augmentation individuelle accordée à tous les salariés indépendemment de leur catégorie socio-professionnelle, en pourcentage ou en euros",
        )
    ]
)

prime = Object(
    id="prime_ppv",
    description = "Information sur la distribution d'une prime de partage de la valeur ajoutée aux salariés, aussi appelé prime macron ou pepa ou PPV. L'unité est toujours l'euro et jamais autre chose",
    examples= [
        (
            """La direction s'accord avec les organisations syndicales (OS) sur le versement d'une prime de 2000 euros à tous les salariés""",
            [{"primes":"2000"}]
        )
    ],
    attributes=[
        Number(
            id="prime_ppv",
            description = "Information sur la distribution d'une prime de partage de la valeur ajoutée aux salariés, aussi appelé prime macron ou pepa ou PPV. L'unité est toujours l'euro et jamais autre chose\
                il ne s'agit pas d'une prime transport ou d'une indemnité. Cette prime ne porte que trois noms: 'partage de la valeur ajoutée', 'pepa', 'macron'",
        )
    ]
)

node = Object(
    id="root_node",
    attributes=[
        augmentation_gen,
        augmentation_gen_ouv,
        augmentation_gen_int,
        augmentation_gen_cad,
        augmentation_ind,
        augmentation_ind_ouv,
        augmentation_ind_int,
        augmentation_ind_cad,
        prime
    ]
)

# create Pydantic model

class Augmentation(BaseModel):
    augm_gen: Optional[float] = Field(..., description="Augmentation générale de salaire/salaire de base pour tous les salariés, en pourcentage ou en euros")
    augm_gen_ouv : Optional[float] = Field(..., description="Augmentation générale de salaire/salaire de base pour les ouvriers et les employés, en pourcentage ou en euros")
    augm_gen_int : Optional[float] = Field(..., description="Augmentation générale de salaire/salaire de base pour les professions intermédiaires, techniciens et agent de maîtrise, en pourcentage ou en euros")
    augm_gen_cad : Optional[float] = Field(..., description="Augmentation générale de salaire/salaire de base pour les cadres et ingénieurs, en pourcentage ou en euros")
    augm_ind: Optional[float] = Field(..., description="Augmentation individuelle de salaire ou mérite pour tous les salariés, en pourcentage ou en euros")
    augm_ind_ouv : Optional[float] = Field(..., description="Augmentation individuelle de salaire ou mérite pour les ouvriers et les employés, en pourcentage ou en euros")
    augm_ind_int : Optional[float] = Field(..., description="Augmentation individuelle de salaire ou mérite pour les professions intermédiaires, techniniciens, en pourcentage ou en euros")
    augm_ind_cad : Optional[float] = Field(..., description="Augmentation individuelle de salaire ou mérite pour les cadres et ingénieurs, en pourcentage ou en euros")
    prime : Optional[int] = Field(...,description="Information sur la distribution d'une prime de partage de la valeur ajoutée aux salariés, aussi appelé prime macron ou pepa ou PPV. L'unité est toujours l'euro et jamais autre chose")
    
    @validator('augm_gen') 
    def no_empty(cls, v):
        if not v:
            raise ValueError('NA')
        return v
    
node, extraction_validator = from_pydantic(
    Augmentation,
    description="Augmentation de la masse salariale générale ou au mérite/individuelle",
    many=True,)

directory = r"C:\Users\garsonj\Desktop\NLP\new_sample_docx"
parent_directory = r"C:\Users\garsonj\Desktop\NLP\vector3"

# Iterate over each file in the directory
for file in glob.glob(os.path.join(directory, '*.docx')):
    file_name = os.path.basename(file)[0:27]
    path = os.path.join(parent_directory, file_name)
    os.makedirs(path, exist_ok=True)
    persist = path

    content = []
    doc = Document()
    loader = Docx2txtLoader(file)
    doc_text = loader.load()
    
    # Retrieve the text from the document
    embedding = OpenAIEmbeddings()
    vectordb = Chroma.from_documents(documents=doc_text, embedding=embedding, persist_directory=persist)
    vectordb.persist()
    vectordb = None

    db = Chroma(persist_directory=persist, embedding_function=embedding)
    

    # Using text-davinci-003 and a temperature of 0
    llm = ChatOpenAI(
    model_name = "gpt-3.5-turbo-16k",
    temperature=0,
    max_tokens=750, #token for completion
    openai_api_key=openai_api_key
    )
    retriever = db.as_retriever()
    qa = RetrievalQA.from_chain_type(llm=llm, chain_type="stuff", retriever=retriever)
    try:
        query = """fais un résumé du texte en relevant les augmentations générales, individuelles et les primes de partage de la valeur ajoutée\
        pour les cadres, intermédiaires et ouvriers, employés"""
        data = qa.run(query)
        print(data)

        content.append(data)
        doc.add_heading(file_name, level=1)
        doc.add_paragraph(data)
        doc.save(file_name + '.docx')
    except:
        print(f'text trop long',{file_name})

def move_files(currrent_directory, directory):
    for file in cd:
        if file.endswith('.docx'):
            shutil.move(file, directory_summary)
    return
    
directory_summary = r"C:\Users\garsonj\Desktop\NLP\summaryGPT4"
cd = glob.glob(r"C:\Users\garsonj\Desktop\NLP\*.docx")
move_files(file, directory_summary)