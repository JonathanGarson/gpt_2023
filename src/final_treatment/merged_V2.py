from docx import Document
import os
import glob

def merge_files(folder_path):
    docx_files = glob.glob(folder_path + '\\*.docx')
    merged_doc = Document()  # Create a new document to store the merged content
    for file in docx_files:
        doc = Document(file)
        for para in doc.paragraphs:
            content = para.text
            style = para.style
            merged_doc.add_paragraph(content, style=style)  # Add paragraph with preserved style

    merged_doc.save('merged.docx')  # Save the merged document


# Provide the folder path where the algorithm will iterate over all *.docx files
folder_path = '.\\sample.docx'
merge_files(folder_path)
